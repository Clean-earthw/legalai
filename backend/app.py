import os
import uuid
import tempfile
import logging
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
import asyncio

from agents.rag.document_store import DocumentStore
from agents.crews.legal_support_agents import LegalSupportAgents, create_legal_support_system

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("legal-support-api")

# Try to import optional dependencies
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyMuPDF not installed. PDF processing will be disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed. DOCX processing will be disabled.")

# Fallback text extraction methods
try:
    import PyPDF2
    PDF_FALLBACK = True
except ImportError:
    PDF_FALLBACK = False

try:
    import docx2txt
    DOCX_FALLBACK = True
except ImportError:
    DOCX_FALLBACK = False

app = FastAPI(
    title="Legal Support API",
    description="API for legal document processing and question answering",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global agents system instance
agents_system = None

# Pydantic models
class QuestionRequest(BaseModel):
    question: str = Field(..., description="The question to answer")
    document_type: Optional[str] = Field(None, description="Filter by document type")

class QuestionResponse(BaseModel):
    answer: str
    sources: List[str] = []
    confidence: Optional[float] = None
    processing_time: float

class UploadResponse(BaseModel):
    document_id: str
    document_name: str
    chunks_added: int
    status: str

class DocumentInfo(BaseModel):
    document_id: str
    document_name: str
    document_type: str
    chunks_count: int

class HealthStatus(BaseModel):
    status: str
    components: Dict[str, str]
    details: Optional[Dict[str, Any]] = None

class SearchRequest(BaseModel):
    query: str
    limit: int = Field(5, ge=1, le=20)
    document_type: Optional[str] = None

class SearchResult(BaseModel):
    text: str
    document_name: str
    document_type: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    similarity_score: Optional[float] = None

class SearchResponse(BaseModel):
    query: str
    results_count: int
    results: List[SearchResult]

# Global agents system instance
agents_system = None

# Pydantic models (keep your existing models)

# Dependency to get agents system
async def get_agents_system():
    global agents_system
    if agents_system is None:
        agents_system = create_legal_support_system(debug_enabled=False)
        # Initialize RAG system
        loop = asyncio.get_event_loop()
        rag_initialized = await loop.run_in_executor(None, agents_system.initialize_rag)
        if not rag_initialized:
            logger.error("Failed to initialize RAG system")
            raise HTTPException(status_code=500, detail="RAG system initialization failed")
    return agents_system

# Startup event
@app.on_event("startup")
async def startup_event():
    global agents_system
    try:
        agents_system = create_legal_support_system(debug_enabled=False)
        # Initialize RAG system synchronously in executor
        loop = asyncio.get_event_loop()
        rag_initialized = await loop.run_in_executor(None, agents_system.initialize_rag)
        
        if not rag_initialized:
            logger.error("RAG system failed to initialize on startup")
        else:
            logger.info("Legal Support API started successfully")
            
    except Exception as e:
        logger.error(f"Startup error: {e}")
        agents_system = None


# Health check endpoint
@app.get("/health", response_model=HealthStatus)
async def health_check(agents: LegalSupportAgents = Depends(get_agents_system)):
    """Check system health"""
    try:
        health = await agents.health_check()
        return health
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Upload PDF/DOCX documents
@app.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    document_name: Optional[str] = None,
    document_type: Optional[str] = None,
    agents: LegalSupportAgents = Depends(get_agents_system)
):
    """Upload and process PDF or DOCX document"""
    try:
        # Validate file type
        content_type = file.content_type
        filename = file.filename.lower()
        
        if not (filename.endswith(('.pdf', '.docx', '.txt')) or 
                content_type in ['application/pdf', 
                               'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                               'text/plain']):
            raise HTTPException(
                status_code=400, 
                detail="Only PDF, DOCX, and TXT files are supported"
            )

        # Generate document name if not provided
        if not document_name:
            document_name = file.filename

        # Auto-detect document type if not provided
        if not document_type:
            if filename.endswith('.pdf'):
                document_type = "pdf"
            elif filename.endswith('.docx'):
                document_type = "docx"
            elif filename.endswith('.txt'):
                document_type = "text"
            else:
                document_type = "unknown"

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        try:
            # Process based on file type
            if filename.endswith('.pdf'):
                result = await process_pdf_file(tmp_file_path, document_name, document_type, agents)
            elif filename.endswith('.docx'):
                result = await process_docx_file(tmp_file_path, document_name, document_type, agents)
            elif filename.endswith('.txt'):
                result = await process_text_file(tmp_file_path, document_name, document_type, agents)
                
            return UploadResponse(
                document_id=result["document_id"],
                document_name=document_name,
                chunks_added=result["chunks_added"],
                status="success"
            )
            
        finally:
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


async def process_pdf_file(file_path: str, document_name: str, document_type: str, agents: LegalSupportAgents):
    """Process PDF file with multiple fallback methods"""
    text_content = []
    
    # Method 1: PyMuPDF (best quality)
    if PDF_SUPPORT:
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text", sort=True)
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text}")
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
            text_content = []
    
    # Method 2: PyPDF2 fallback
    if not text_content and PDF_FALLBACK:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"Page {page_num + 1}:\n{text}")
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
    
    if not text_content:
        raise HTTPException(
            status_code=400, 
            detail="Could not extract text from PDF. Please ensure the PDF contains selectable text."
        )
    
    full_text = "\n\n".join(text_content)
    
    # Add to document store - FIXED PARAMETERS
    loop = asyncio.get_event_loop()
    try:
        result = await loop.run_in_executor(
            None, 
            agents.add_document_from_text,
            full_text,          # text
            document_name,      # document_name
            None,              # document_id
            document_type,     # document_type (not splitter_type!)
            "recursive"        # splitter_type
        )
        return result
    except Exception as e:
        logger.error(f"Error adding document to store: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to add document to knowledge base: {str(e)}")


async def process_docx_file(file_path: str, document_name: str, document_type: str, agents: LegalSupportAgents):
    """Process DOCX file with multiple fallback methods"""
    text_content = []
    
    # Method 1: python-docx
    if DOCX_SUPPORT:
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
                            
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
            text_content = []
    
    # Method 2: docx2txt fallback
    if not text_content and DOCX_FALLBACK:
        try:
            text = docx2txt.process(file_path)
            if text.strip():
                text_content = [text]
        except Exception as e:
            logger.warning(f"docx2txt extraction failed: {e}")
    
    if not text_content:
        raise HTTPException(
            status_code=400, 
            detail="Could not extract text from DOCX file."
        )
    
    full_text = "\n".join(text_content)
    
    # Add to document store - FIXED PARAMETERS
    loop = asyncio.get_event_loop()
    try:
        result = await loop.run_in_executor(
            None, 
            agents.add_document_from_text,
            full_text,          # text
            document_name,      # document_name
            None,              # document_id
            document_type,     # document_type
            "recursive"        # splitter_type
        )
        return result
    except Exception as e:
        logger.error(f"Error adding document to store: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to add document to knowledge base: {str(e)}")


async def process_text_file(file_path: str, document_name: str, document_type: str, agents: LegalSupportAgents):
    """Process TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        # Add to document store - FIXED PARAMETERS
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None, 
            agents.add_document_from_text,
            text_content,       # text
            document_name,      # document_name
            None,              # document_id
            document_type,     # document_type
            "recursive"        # splitter_type
        )
        return result
    except Exception as e:
        logger.error(f"Error processing text file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process text file: {str(e)}")

# Ulpoad text content directly - FIXED
@app.post("/upload/text", response_model=UploadResponse)
async def upload_text(
    text: str,
    document_name: str,
    document_type: str = "text",
    agents: LegalSupportAgents = Depends(get_agents_system)
):
    """Upload text content directly"""
    try:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            agents.add_document_from_text,
            text,               # text
            document_name,      # document_name  
            None,              # document_id
            document_type,     # document_type
            "recursive"        # splitter_type
        )
        
        return UploadResponse(
            document_id=result["document_id"],
            document_name=document_name,
            chunks_added=result["chunks_added"],
            status="success"
        )
        
    except Exception as e:
        logger.error(f"Text upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
# Ask question endpoint
@app.post("/ask", response_model=QuestionResponse)
async def ask_question(
    request: QuestionRequest,
    agents: LegalSupportAgents = Depends(get_agents_system)
):
    """Ask a question based on uploaded documents"""
    try:
        import time
        start_time = time.time()
        
        answer = await agents.process_query(request.question)
        
        processing_time = time.time() - start_time
        
        # Extract sources from answer if available
        sources = []
        if "Sources:" in answer:
            parts = answer.split("Sources:")
            answer_text = parts[0].strip()
            sources_text = parts[1].strip(' *')
            sources = [s.strip() for s in sources_text.split(',')]
        else:
            answer_text = answer
        
        return QuestionResponse(
            answer=answer_text,
            sources=sources,
            confidence=0.8,  # You might want to extract this from the actual response
            processing_time=round(processing_time, 2)
        )
        
    except Exception as e:
        logger.error(f"Question processing error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Search documents endpoint
@app.post("/search", response_model=SearchResponse)
async def search_documents(
    request: SearchRequest,
    agents: LegalSupportAgents = Depends(get_agents_system)
):
    """Search for relevant document chunks"""
    try:
        results = agents.search_documents(
            query=request.query,
            limit=request.limit,
            document_type=request.document_type
        )
        
        search_results = []
        for result in results:
            search_results.append(SearchResult(
                text=result.get('text', ''),
                document_name=result.get('document_name', 'Unknown'),
                document_type=result.get('document_type', 'text'),
                page_number=result.get('page_number'),
                section=result.get('section'),
                similarity_score=result.get('_distance')
            ))
        
        return SearchResponse(
            query=request.query,
            results_count=len(search_results),
            results=search_results
        )
        
    except Exception as e:
        logger.error(f"Search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get all documents endpoint
@app.get("/documents", response_model=List[DocumentInfo])
async def get_documents(agents: LegalSupportAgents = Depends(get_agents_system)):
    """Get list of all uploaded documents"""
    try:
        documents = agents.get_all_documents()
        return [
            DocumentInfo(
                document_id=doc['document_id'],
                document_name=doc['document_name'],
                document_type=doc.get('document_type', 'text'),
                chunks_count=doc.get('chunks_count', 0)
            )
            for doc in documents
        ]
        
    except Exception as e:
        logger.error(f"Get documents error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Delete document endpoint
@app.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    agents: LegalSupportAgents = Depends(get_agents_system)
):
    """Delete a document and all its chunks"""
    try:
        result = agents.delete_document(document_id)
        return {
            "document_id": document_id,
            "chunks_deleted": result.get("chunks_deleted", 0),
            "status": "deleted"
        }
        
    except Exception as e:
        logger.error(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "Legal Support API",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "upload": "/upload",
            "ask": "/ask",
            "search": "/search",
            "documents": "/documents"
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)