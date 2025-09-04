import os
import uuid
import tempfile
import logging
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Query, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field

from agents.crews.legal_support_agents import UserLegalSupportAgents, create_user_legal_support_system

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("user-legal-support-api")

# Set required environment variable
os.environ["JINA_AI_API_KEY"] = os.getenv("JINA_API_KEY", "")

# Try to import optional dependencies
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyMuPDF not installed. PDF processing will be disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed. DOCX processing will be disabled.")

# Fallback text extraction methods
try:
    import PyPDF2
    PDF_FALLBACK = True
except ImportError:
    PDF_FALLBACK = False

try:
    import docx2txt
    DOCX_FALLBACK = True
except ImportError:
    DOCX_FALLBACK = False

app = FastAPI(
    title="User Legal Support API",
    description="API for user-specific legal document processing and question answering with session management",
    version="2.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global agents system instance
agents_system = None

# Pydantic models
class QuestionRequest(BaseModel):
    question: str = Field(..., description="The question to answer")
    document_type: Optional[str] = Field(None, description="Filter by document type")

class QuestionResponse(BaseModel):
    answer: str
    sources: List[str] = []
    confidence: Optional[float] = None
    processing_time: float
    user_id: str
    session_id: Optional[str] = None

class UploadResponse(BaseModel):
    document_id: str
    document_name: str
    chunks_added: int
    status: str
    user_id: str
    session_id: Optional[str] = None

class DocumentInfo(BaseModel):
    document_id: str
    document_name: str
    document_type: str
    chunks_count: int
    user_id: str
    session_id: Optional[str] = None
    created_at: Optional[str] = None
    file_size: Optional[int] = None

class UserStats(BaseModel):
    user_id: str
    total_documents: int
    total_chunks: int
    sessions: List[str]
    document_types: Dict[str, int]
    total_file_size: int

class HealthStatus(BaseModel):
    status: str
    components: Dict[str, str]
    details: Optional[Dict[str, Any]] = None

class SearchRequest(BaseModel):
    query: str
    limit: int = Field(5, ge=1, le=20)
    document_type: Optional[str] = None

class SearchResult(BaseModel):
    text: str
    document_name: str
    document_type: str
    document_id: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    similarity_score: Optional[float] = None
    user_id: str
    session_id: Optional[str] = None

class SearchResponse(BaseModel):
    query: str
    results_count: int
    results: List[SearchResult]
    user_id: str
    session_id: Optional[str] = None

class SessionRequest(BaseModel):
    session_name: Optional[str] = Field(None, description="Optional session name")

class SessionResponse(BaseModel):
    session_id: str
    user_id: str
    status: str

class DeleteResponse(BaseModel):
    status: str
    chunks_deleted: int
    user_id: str
    document_id: Optional[str] = None
    session_id: Optional[str] = None

class TextUploadRequest(BaseModel):
    text: str = Field(..., description="Text content to upload")
    document_name: str = Field(..., description="Name for the document")
    document_type: str = Field("text", description="Type of document")

# Helper functions for user identification
def get_user_id(x_user_id: Optional[str] = Header(None)) -> str:
    """Extract user ID from headers"""
    if not x_user_id or not x_user_id.strip():
        raise HTTPException(
            status_code=400, 
            detail="User ID is required. Please provide X-User-Id header."
        )
    return x_user_id.strip()

def get_session_id(x_session_id: Optional[str] = Header(None)) -> Optional[str]:
    """Extract optional session ID from headers"""
    if x_session_id and x_session_id.strip():
        return x_session_id.strip()
    return None

# Dependency to get agents system
def get_agents_system():
    global agents_system
    if agents_system is None:
        try:
            agents_system = create_user_legal_support_system(debug_enabled=False)
            # Initialize RAG system
            rag_initialized = agents_system.initialize_rag()
            if not rag_initialized:
                logger.error("Failed to initialize RAG system")
                raise HTTPException(status_code=500, detail="RAG system initialization failed")
        except Exception as e:
            logger.error(f"Error creating agents system: {e}")
            raise HTTPException(status_code=500, detail="System initialization failed")
    return agents_system

# Startup event
@app.on_event("startup")
async def startup_event():
    global agents_system
    try:
        logger.info("Starting User Legal Support API...")
        agents_system = create_user_legal_support_system(debug_enabled=False)
        
        # Initialize RAG system
        rag_initialized = agents_system.initialize_rag()
        
        if not rag_initialized:
            logger.error("RAG system failed to initialize on startup")
        else:
            logger.info("User Legal Support API started successfully")
            
    except Exception as e:
        logger.error(f"Startup error: {e}")
        agents_system = None

# Health check endpoint
@app.get("/health", response_model=HealthStatus)
async def health_check(agents: UserLegalSupportAgents = Depends(get_agents_system)):
    """Check system health"""
    try:
        health = agents.health_check()
        return HealthStatus(
            status=health["status"],
            components=health["components"],
            details=health.get("document_store_details")
        )
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Create new session
@app.post("/sessions", response_model=SessionResponse)
async def create_session(
    request: SessionRequest = SessionRequest(),
    user_id: str = Depends(get_user_id)
):
    """Create a new session for the user"""
    try:
        session_id = str(uuid.uuid4())
        logger.info(f"Created new session {session_id} for user {user_id}")
        
        return SessionResponse(
            session_id=session_id,
            user_id=user_id,
            status="created"
        )
        
    except Exception as e:
        logger.error(f"Session creation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Upload PDF/DOCX documents
@app.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    document_name: Optional[str] = None,
    document_type: Optional[str] = None,
    user_id: str = Depends(get_user_id),
    session_id: Optional[str] = Depends(get_session_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Upload and process PDF or DOCX document for a specific user"""
    try:
        # Validate file type
        content_type = file.content_type
        filename = file.filename.lower() if file.filename else ""
        
        if not (filename.endswith(('.pdf', '.docx', '.txt')) or 
                content_type in ['application/pdf', 
                               'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                               'text/plain']):
            raise HTTPException(
                status_code=400, 
                detail="Only PDF, DOCX, and TXT files are supported"
            )

        # Generate document name if not provided
        if not document_name:
            document_name = file.filename if file.filename else "untitled_document"

        # Auto-detect document type if not provided
        if not document_type:
            if filename.endswith('.pdf'):
                document_type = "pdf"
            elif filename.endswith('.docx'):
                document_type = "docx"
            elif filename.endswith('.txt'):
                document_type = "text"
            else:
                document_type = "unknown"

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1] if filename else '.tmp') as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        try:
            # Process based on file type
            if filename.endswith('.pdf'):
                result = process_pdf_file(tmp_file_path, document_name, document_type, user_id, session_id, agents)
            elif filename.endswith('.docx'):
                result = process_docx_file(tmp_file_path, document_name, document_type, user_id, session_id, agents)
            elif filename.endswith('.txt'):
                result = process_text_file(tmp_file_path, document_name, document_type, user_id, session_id, agents)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type")
                
            return UploadResponse(
                document_id=result["document_id"],
                document_name=document_name,
                chunks_added=result["chunks_added"],
                status="success",
                user_id=user_id,
                session_id=session_id
            )
            
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

def process_pdf_file(file_path: str, document_name: str, document_type: str, 
                    user_id: str, session_id: Optional[str], agents: UserLegalSupportAgents):
    """Process PDF file with multiple fallback methods"""
    text_content = []
    
    # Method 1: PyMuPDF (best quality)
    if PDF_SUPPORT:
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text", sort=True)
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text}")
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
            text_content = []
    
    # Method 2: PyPDF2 fallback
    if not text_content and PDF_FALLBACK:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"Page {page_num + 1}:\n{text}")
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
    
    if not text_content:
        raise HTTPException(
            status_code=400, 
            detail="Could not extract text from PDF. Please ensure the PDF contains selectable text."
        )
    
    full_text = "\n\n".join(text_content)
    
    # Add to user document store
    try:
        result = agents.add_user_document_from_text(
            user_id=user_id,
            text=full_text,
            document_name=document_name,
            session_id=session_id,
            document_id=None,
            document_type=document_type,
            splitter_type="recursive"
        )
        return result
    except Exception as e:
        logger.error(f"Error adding document to store for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to add document to knowledge base: {str(e)}")

def process_docx_file(file_path: str, document_name: str, document_type: str, 
                     user_id: str, session_id: Optional[str], agents: UserLegalSupportAgents):
    """Process DOCX file with multiple fallback methods"""
    text_content = []
    
    # Method 1: python-docx
    if DOCX_SUPPORT:
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
                            
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
            text_content = []
    
    # Method 2: docx2txt fallback
    if not text_content and DOCX_FALLBACK:
        try:
            text = docx2txt.process(file_path)
            if text.strip():
                text_content = [text]
        except Exception as e:
            logger.warning(f"docx2txt extraction failed: {e}")
    
    if not text_content:
        raise HTTPException(
            status_code=400, 
            detail="Could not extract text from DOCX file."
        )
    
    full_text = "\n".join(text_content)
    
    # Add to user document store
    try:
        result = agents.add_user_document_from_text(
            user_id=user_id,
            text=full_text,
            document_name=document_name,
            session_id=session_id,
            document_id=None,
            document_type=document_type,
            splitter_type="recursive"
        )
        return result
    except Exception as e:
        logger.error(f"Error adding document to store for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to add document to knowledge base: {str(e)}")

def process_text_file(file_path: str, document_name: str, document_type: str, 
                     user_id: str, session_id: Optional[str], agents: UserLegalSupportAgents):
    """Process TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="Text file is empty")
        
        # Add to user document store
        result = agents.add_user_document_from_text(
            user_id=user_id,
            text=text_content,
            document_name=document_name,
            session_id=session_id,
            document_id=None,
            document_type=document_type,
            splitter_type="recursive"
        )
        return result
    except Exception as e:
        logger.error(f"Error processing text file for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process text file: {str(e)}")

# Upload text content directly
@app.post("/upload/text", response_model=UploadResponse)
async def upload_text(
    request: TextUploadRequest,
    user_id: str = Depends(get_user_id),
    session_id: Optional[str] = Depends(get_session_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Upload text content directly for a specific user"""
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text content cannot be empty")
            
        result = agents.add_user_document_from_text(
            user_id=user_id,
            text=request.text,
            document_name=request.document_name,
            session_id=session_id,
            document_id=None,
            document_type=request.document_type,
            splitter_type="recursive"
        )
        
        return UploadResponse(
            document_id=result["document_id"],
            document_name=request.document_name,
            chunks_added=result["chunks_added"],
            status="success",
            user_id=user_id,
            session_id=session_id
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Text upload error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Ask question endpoint
@app.post("/ask", response_model=QuestionResponse)
async def ask_question(
    request: QuestionRequest,
    user_id: str = Depends(get_user_id),
    session_id: Optional[str] = Depends(get_session_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Ask a question based on user's uploaded documents"""
    try:
        import time
        start_time = time.time()
        
        answer = agents.process_query(user_id, request.question, session_id)
        
        processing_time = time.time() - start_time
        
        # Extract sources from answer if available
        sources = []
        answer_text = answer
        if "*Sources from your documents:" in answer:
            parts = answer.split("*Sources from your documents:")
            answer_text = parts[0].strip()
            if len(parts) > 1:
                sources_text = parts[1].strip(' *')
                sources = [s.strip() for s in sources_text.split(',') if s.strip()]
        
        return QuestionResponse(
            answer=answer_text,
            sources=sources,
            confidence=0.8,
            processing_time=round(processing_time, 2),
            user_id=user_id,
            session_id=session_id
        )
        
    except Exception as e:
        logger.error(f"Question processing error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Search documents endpoint
@app.post("/search", response_model=SearchResponse)
async def search_documents(
    request: SearchRequest,
    user_id: str = Depends(get_user_id),
    session_id: Optional[str] = Depends(get_session_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Search for relevant document chunks for a specific user"""
    try:
        results = agents.search_user_documents(
            user_id=user_id,
            query=request.query,
            limit=request.limit,
            session_id=session_id,
            document_type=request.document_type
        )
        
        search_results = []
        for result in results:
            search_results.append(SearchResult(
                text=result.get('text', ''),
                document_name=result.get('document_name', 'Unknown'),
                document_type=result.get('document_type', 'text'),
                document_id=result.get('document_id', ''),
                page_number=result.get('page_number'),
                section=result.get('section'),
                similarity_score=result.get('_distance'),
                user_id=result.get('user_id', user_id),
                session_id=result.get('session_id')
            ))
        
        return SearchResponse(
            query=request.query,
            results_count=len(search_results),
            results=search_results,
            user_id=user_id,
            session_id=session_id
        )
        
    except Exception as e:
        logger.error(f"Search error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))
# Get documents endpoint - User-specific implementation
@app.get("/documents", response_model=List[DocumentInfo])
async def get_documents(
    user_id: str = Depends(get_user_id),
    session_id: Optional[str] = Depends(get_session_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Get list of documents for a specific user"""
    try:
        # This correctly uses get_user_documents for user-specific document retrieval
        # NOT get_all_system_documents which would return documents from all users
        documents = agents.get_user_documents(user_id, session_id)
        
        # Process and return user-specific documents
        document_list = []
        for doc in documents:
            document_info = DocumentInfo(
                document_id=doc['document_id'],
                document_name=doc['document_name'],
                document_type=doc.get('document_type', 'text'),
                user_id=doc.get('user_id', user_id),
                session_id=doc.get('session_id'),
                created_at=doc.get('created_at'),
                file_size=doc.get('file_size'),
                chunks_count=doc.get('chunks_count', 0)
            )
            document_list.append(document_info)
        
        return document_list
        
    except Exception as e:
        logger.error(f"Get documents error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# If you need a separate admin endpoint to get all system documents, 
# you could add this as a separate endpoint with proper authorization:
@app.get("/admin/documents", response_model=List[DocumentInfo])
async def get_all_system_documents_admin(
    # Add proper admin authentication here
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Admin endpoint to get all documents from all users - requires admin access"""
    try:
        # Only use this in admin contexts with proper authentication
        documents = agents.get_all_documents()
        
        document_list = []
        for doc in documents:
            document_info = DocumentInfo(
                document_id=doc['document_id'],
                document_name=doc['document_name'],
                document_type=doc.get('document_type', 'text'),
                user_id=doc.get('user_id', ''),
                session_id=doc.get('session_id'),
                created_at=doc.get('created_at'),
                file_size=doc.get('file_size'),
                chunks_count=doc.get('chunks_count', 0)
            )
            document_list.append(document_info)
        
        return document_list
        
    except Exception as e:
        logger.error(f"Get all system documents error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get user statistics
@app.get("/stats", response_model=UserStats)
async def get_user_stats(
    user_id: str = Depends(get_user_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Get statistics for a specific user's documents"""
    try:
        stats = agents.get_user_stats(user_id)
        return UserStats(
            user_id=stats['user_id'],
            total_documents=stats.get('total_documents', 0),
            total_chunks=stats.get('total_chunks', 0),
            sessions=stats.get('sessions', []),
            document_types=stats.get('document_types', {}),
            total_file_size=stats.get('total_file_size', 0)
        )
        
    except Exception as e:
        logger.error(f"Get stats error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Delete document endpoint
@app.delete("/documents/{document_id}", response_model=DeleteResponse)
async def delete_document(
    document_id: str,
    user_id: str = Depends(get_user_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Delete a document and all its chunks for a specific user"""
    try:
        result = agents.delete_user_document(user_id, document_id)
        return DeleteResponse(
            status="deleted",
            chunks_deleted=result.get("chunks_deleted", 0),
            user_id=user_id,
            document_id=document_id
        )
        
    except Exception as e:
        logger.error(f"Delete document error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Delete session documents endpoint
@app.delete("/sessions/{session_id}/documents", response_model=DeleteResponse)
async def delete_session_documents(
    session_id: str,
    user_id: str = Depends(get_user_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Delete all documents for a specific user session"""
    try:
        result = agents.delete_user_session_documents(user_id, session_id)
        return DeleteResponse(
            status="deleted",
            chunks_deleted=result.get("chunks_deleted", 0),
            user_id=user_id,
            session_id=session_id
        )
        
    except Exception as e:
        logger.error(f"Delete session documents error for user {user_id}, session {session_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Delete all user documents endpoint
@app.delete("/users/{user_id}/documents", response_model=DeleteResponse)
async def delete_all_user_documents(
    user_id_path: str,
    user_id: str = Depends(get_user_id),
    agents: UserLegalSupportAgents = Depends(get_agents_system)
):
    """Delete all documents for a specific user across all sessions"""
    try:
        # Ensure the path user_id matches the header user_id for security
        if user_id_path != user_id:
            raise HTTPException(status_code=403, detail="User ID mismatch")
            
        result = agents.delete_all_user_documents(user_id)
        return DeleteResponse(
            status="deleted",
            chunks_deleted=result.get("chunks_deleted", 0),
            user_id=user_id
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete all user documents error for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "User Legal Support API",
        "version": "2.0.0",
        "description": "Multi-user legal document processing and AI-powered question answering with session management",
        "endpoints": {
            "health": "/health",
            "create_session": "POST /sessions",
            "upload": "POST /upload",
            "upload_text": "POST /upload/text", 
            "ask": "POST /ask",
            "search": "POST /search",
            "documents": "GET /documents",
            "stats": "GET /stats",
            "delete_document": "DELETE /documents/{document_id}",
            "delete_session": "DELETE /sessions/{session_id}/documents",
            "delete_all_user_docs": "DELETE /users/{user_id}/documents"
        },
        "headers_required": {
            "X-User-Id": "Required for all endpoints (except health and root)",
            "X-Session-Id": "Optional session identifier for document isolation"
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)